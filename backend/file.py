from PyPDF2 import PdfReader
from langchain.text_splitter import TokenTextSplitter
from langchain_community.vectorstores import FAISS
import os
from dotenv import load_dotenv
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import docx
from langchain.docstore.document import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Preformatted
from reportlab.lib.styles import getSampleStyleSheet
from bs4 import BeautifulSoup

load_dotenv()

embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

def limpar_texto(texto):
    return " ".join(texto.split())

def process_document(files, file_extension):
    if file_extension == 'pdf':
        print("PROCESSANDO PDF")
        reader = PdfReader(files)
        raw_text = ''
        for page in reader.pages:
            text = page.extract_text()
            if text:
                raw_text += text
    elif file_extension == 'txt':
        print("PROCESSANDO TXT")
        if isinstance(files, BytesIO):
            files.seek(0)
            raw_text = files.read().decode('utf-8')
        else:
            with open(files, 'r') as file:
                raw_text = file.read()
    elif file_extension == 'docx':
        print("PROCESSANDO DOCX")
        doc = docx.Document(files)
        raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file type")
    
    text_splitter = TokenTextSplitter(
        chunk_size=350,
        chunk_overlap=100
    )
    
    documents = text_splitter.split_text(raw_text)
    document_objs = [Document(page_content=doc) for doc in documents]
    return document_objs

def criar_indices_faiss(files, file_extension):
    documents = process_document(files, file_extension)
    # Cria o índice FAISS a partir dos textos e dos embeddings
    vector_store = FAISS.from_documents(documents, embeddings)
    # Salva o índice em um arquivo
    vector_store.save_local("faiss_index")
    return vector_store

def carregar_indices_faiss():
    # Carrega o índice FAISS a partir do arquivo
    vector_store = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    return vector_store

def adicionar_texto_ao_indice(vector_store, files, file_extension):
    documents = process_document(files, file_extension)
    # Adiciona novos textos ao índice existente
    vector_store.add_documents(documents)
    # Salva o índice atualizado em um arquivo
    vector_store.save_local("faiss_index")

def verificar_e_atualizar_indice(files, file_extension):
    if os.path.exists("faiss_index"):
        vector_store = carregar_indices_faiss()
        adicionar_texto_ao_indice(vector_store, files, file_extension)
    else:
        vector_store = criar_indices_faiss(files, file_extension)
    return vector_store

def procurar_similaridade(query):
    print('QUERY:', query)
    # Verificar e atualizar o índice
    vector_store = carregar_indices_faiss()
    print("Vector store carregado")
    # Realiza a busca de similaridade
    resultados = vector_store.similarity_search(query=query, k=3)
    print("RESULTADOS OBTIDOS")
    print(resultados)
    textos_resultados = [limpar_texto(doc.page_content) for doc in resultados]
    textos_resultados = ''.join(textos_resultados)
    return textos_resultados

# Função para criar um PDF com a resposta da API
def criar_pdf(resposta_da_api, diretorio):

    print(resposta_da_api)
    if len(resposta_da_api) >= 20:
            nome_do_arquivo = resposta_da_api[:20] + '.pdf'
    else:
            nome_do_arquivo = resposta_da_api + '.pdf'
    
    nome_do_arquivo = remove_html_tags(nome_do_arquivo)

    if diretorio:
        # Verifica se o caminho é absoluto ou relativo
        if not os.path.isabs(diretorio):
            # Converte o caminho relativo para absoluto
            diretorio = os.path.join(os.getcwd(), diretorio)
        caminho_completo = os.path.join(diretorio, nome_do_arquivo)
    else:
        caminho_completo = nome_do_arquivo
    
    # Cria o diretório se não existir
    if diretorio and not os.path.exists(diretorio):
        os.makedirs(diretorio)
    
    # Inicia um novo documento PDF
    doc = SimpleDocTemplate(caminho_completo, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = []

     # Adiciona a resposta da API ao PDF como um bloco preformatado
    preformatted_text = Preformatted(resposta_da_api, styles['Normal'])
    flowables.append(preformatted_text)
    
    # Salva o PDF
    doc.build(flowables)

def remove_html_tags(text):
    soup = BeautifulSoup(text, "html.parser")
    return soup.get_text()
